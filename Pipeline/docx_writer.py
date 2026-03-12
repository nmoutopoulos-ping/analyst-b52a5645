"""
Word document generation for investment summaries.

This module generates professional multi-page Word (.docx) documents that contain
investment analysis summaries, including property details, rental comparables analysis,
and projected income information.
"""

from datetime import datetime
from collections import defaultdict
from pathlib import Path

from helpers import shorten_address, unit_counts as _unit_counts


def generate_summary_docx(job_dir, search_meta, comp_summary, all_comp_rows=None):
    """
    Generate a professional multi-page investment summary as a Word .docx file.
    Page 1: Executive summary — property details, aggregated comp analysis, projected income.
    Page 2+: Full comp listings table, grouped by unit type, with address, rent, SF, dist.
    Saved to job_dir as Ping_{searchId}_Summary.docx.
    Returns the Path to the saved file.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Colour palette ─────────────────────────────────────────
    NAVY       = "0F172A"
    LIGHT_NAVY = "1E3A5F"
    SLATE_BG   = "F8FAFC"
    BLUE_BG    = "EFF6FF"

    # ── XML helpers ────────────────────────────────────────────
    def _shading(cell_or_tc, hex_color):
        tc = getattr(cell_or_tc, "_tc", cell_or_tc)
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _cell_margins(cell, top=60, bottom=60, left=120, right=120):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        # OOXML schema order: top → start → bottom → end
        for side, val in [("top", top), ("start", left),
                          ("bottom", bottom), ("end", right)]:
            e = OxmlElement(f"w:{side}")
            e.set(qn("w:w"), str(val))
            e.set(qn("w:type"), "dxa")
            tcMar.append(e)
        tcPr.append(tcMar)

    def _tbl_borders(table, val="none", color="auto", sz="0"):
        tbl  = table._tbl
        tblPr = tbl.tblPr
        bdr  = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"),   val)
            b.set(qn("w:sz"),    sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            bdr.append(b)
        # Insert after tblW so element order matches OOXML schema:
        # tblW → tblBorders → tblCellMar / tblLook
        tblW_el = tblPr.find(qn("w:tblW"))
        if tblW_el is not None:
            pos = list(tblPr).index(tblW_el) + 1
        else:
            pos = 0
        tblPr.insert(pos, bdr)

    def _light_borders(table):
        _tbl_borders(table, val="single", color="E2E8F0", sz="4")

    def _para_border_bottom(para, color="3B82F6", sz="6"):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b    = OxmlElement("w:bottom")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "2")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        # Insert before spacing/jc/rPr so element order matches OOXML schema
        pPr.insert(0, pBdr)

    def _run(para, text, bold=False, size=9, color="0F172A",
             italic=False, align=None):
        if align:
            para.alignment = align
        r = para.add_run(text)
        r.bold      = bold
        r.italic    = italic
        r.font.size = Pt(size)
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16),
        )
        return r

    def _section_header(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        _para_border_bottom(p)
        _run(p, text, bold=True, size=9, color="1E3A5F")
        return p

    # ── Pre-compute numbers ────────────────────────────────────
    short_name  = shorten_address(search_meta["address"])
    date_str    = datetime.now().strftime("%B %d, %Y")

    try:    price       = float(search_meta.get("price") or 0)
    except: price = 0
    try:    cost        = float(search_meta.get("cost")  or 0)
    except: cost = 0
    try:    sqft        = float(search_meta.get("sqft")  or 0)
    except: sqft = 0
    try:    total_units = int(float(search_meta.get("totalUnits") or 0))
    except: total_units = 0

    n_combos        = len(comp_summary)
    unit_cnts       = _unit_counts(search_meta["totalUnits"], n_combos)
    gross_monthly   = sum(
        unit_cnts[i] * (s.get("avg_rent") or 0)
        for i, s in enumerate(comp_summary)
    )
    gross_annual = gross_monthly * 12

    # ── CoC estimates using template's fixed underwriting assumptions ──────────
    # These mirror the hardcoded defaults in the Excel template
    _LTV          = 0.70
    _CLOSING_PCT  = 0.02
    _VACANCY      = 0.07
    _OTHER_INC_MO = 75        # $ / unit / month
    _OPEX_RATIO   = 0.35      # % of EGI
    _INT_RATE     = 0.065     # annual (IO in Year 1)
    _RENT_GROWTH1 = 0.03      # Year-1 rent growth applied by the model

    _gpr1   = gross_annual * (1 + _RENT_GROWTH1)
    _egi1   = _gpr1 * (1 - _VACANCY) + (_OTHER_INC_MO * (total_units or 1) * 12)
    _noi1   = _egi1 * (1 - _OPEX_RATIO)
    _all_in = price * (1 + _CLOSING_PCT) + cost          # total capital deployed
    _equity = price * (1 - _LTV + _CLOSING_PCT) + cost   # equity invested
    _loan   = price * _LTV
    _ds1    = _loan * _INT_RATE                           # Year-1 IO debt service

    _unlev_coc = (_noi1 / _all_in)   if _all_in  else 0
    _lev_coc   = ((_noi1 - _ds1) / _equity) if _equity else 0

    # ── Build document ─────────────────────────────────────────
    doc  = Document()
    sect = doc.sections[0]
    sect.page_width    = Inches(8.5)
    sect.page_height   = Inches(11)
    sect.left_margin   = Inches(0.75)
    sect.right_margin  = Inches(0.75)
    sect.top_margin    = Inches(0.75)
    sect.bottom_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name  = "Arial"
    normal.font.size  = Pt(9)
    normal.paragraph_format.space_after  = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    CW = Inches(7.0)   # content width (8.5 - 2×0.75)

    # ── HEADER: dark navy bar ─────────────────────────────────
    hdr = doc.add_table(rows=1, cols=2)
    hdr.width = CW
    _tbl_borders(hdr)
    lc, rc = hdr.rows[0].cells[0], hdr.rows[0].cells[1]
    lc.width = Inches(5.0)
    rc.width = Inches(2.0)
    for c in (lc, rc):
        _shading(c, NAVY)
        _cell_margins(c, top=120, bottom=120, left=160, right=160)

    _run(lc.paragraphs[0], short_name.upper(),
         bold=True, size=13, color="FFFFFF")

    ping_p = rc.paragraphs[0]
    ping_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(ping_p, "PING",               bold=True, size=18, color="3B82F6")
    _run(ping_p, "\nInvestment Analysis", size=7.5, color="94A3B8")

    # ── METADATA BAR: navy strip ──────────────────────────────
    meta_items = [
        ("SEARCH ID",      search_meta.get("searchId", "—")),
        ("DATE",           date_str),
        ("RADIUS",         f"{search_meta.get('radius', '—')} mi"),
        ("STATUS FILTER",  search_meta.get("status", "—")),
    ]
    meta = doc.add_table(rows=1, cols=4)
    meta.width = CW
    _tbl_borders(meta)
    for i, (lbl, val) in enumerate(meta_items):
        c = meta.rows[0].cells[i]
        _shading(c, LIGHT_NAVY)
        _cell_margins(c, top=80, bottom=80, left=120, right=80)
        p = c.paragraphs[0]
        _run(p, lbl + "\n", bold=True, size=6.5, color="94A3B8")
        _run(p, val,        bold=True, size=9,   color="FFFFFF")

    # ── SUBJECT PROPERTY ──────────────────────────────────────
    _section_header(doc, "SUBJECT PROPERTY")
    prop_rows = [
        ("Acquisition Price",      f"${price:,.0f}"               if price       else "—"),
        ("Estimated Improvements", f"${cost:,.0f}"                if cost        else "—"),
        ("All-In Cost",            f"${price + cost:,.0f}"        if price       else "—"),
        ("Total Units",            str(total_units)               if total_units else "—"),
        ("Building SF",            f"{sqft:,.0f}"                 if sqft        else "—"),
        ("Price / Unit",           f"${price/total_units:,.0f}"   if price and total_units else "—"),
    ]
    prop = doc.add_table(rows=len(prop_rows), cols=2)
    prop.width = CW
    _light_borders(prop)
    for i, (lbl, val) in enumerate(prop_rows):
        bg = SLATE_BG if i % 2 == 0 else "FFFFFF"
        lc, vc = prop.rows[i].cells[0], prop.rows[i].cells[1]
        lc.width = Inches(3.5);  vc.width = Inches(3.5)
        _shading(lc, bg);        _shading(vc, bg)
        _cell_margins(lc);       _cell_margins(vc)
        _run(lc.paragraphs[0], lbl,  size=9, color="475569")
        _run(vc.paragraphs[0], val,  size=9, color="0F172A", bold=True)

    # ── RENTAL COMP ANALYSIS ──────────────────────────────────
    _section_header(doc, "RENTAL COMP ANALYSIS")
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(0)
    note_p.paragraph_format.space_after  = Pt(4)
    _run(note_p,
         f"Live market data  ·  {search_meta.get('radius', '—')} mi radius  "
         f"·  {search_meta.get('status', '—')} listings",
         size=8, color="64748B", italic=True)

    comp_hdrs  = ["Unit Type", "Comps", "Avg Rent/Mo", "Avg SF", "Rent/SF"]
    comp_widths = [Inches(1.6), Inches(0.8), Inches(1.7), Inches(1.5), Inches(1.4)]

    ctbl = doc.add_table(rows=1 + len(comp_summary), cols=5)
    ctbl.width = CW
    _light_borders(ctbl)

    for i, (h, w) in enumerate(zip(comp_hdrs, comp_widths)):
        c = ctbl.rows[0].cells[i]
        c.width = w
        _shading(c, NAVY)
        _cell_margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size=8.5, color="FFFFFF")

    for idx, s in enumerate(comp_summary):
        bg       = SLATE_BG if idx % 2 == 0 else "FFFFFF"
        avg_rent = s.get("avg_rent") or 0
        avg_sqft = s.get("avg_sqft") or 0
        rent_psf = avg_rent / avg_sqft if avg_sqft else 0
        vals = [
            (f"{s['beds']}BD/{s['baths']}BA",  WD_ALIGN_PARAGRAPH.LEFT,   False),
            (str(s.get("count", 0)),           WD_ALIGN_PARAGRAPH.CENTER, False),
            (f"${avg_rent:,.0f}" if avg_rent else "—",
                                               WD_ALIGN_PARAGRAPH.CENTER, True),
            (f"{avg_sqft:,.0f}" if avg_sqft else "—",
                                               WD_ALIGN_PARAGRAPH.CENTER, False),
            (f"${rent_psf:.2f}" if rent_psf else "—",
                                               WD_ALIGN_PARAGRAPH.CENTER, True),
        ]
        row = ctbl.rows[idx + 1]
        for i, (val, align, bold) in enumerate(vals):
            c = row.cells[i]
            c.width = comp_widths[i]
            _shading(c, bg)
            _cell_margins(c)
            _run(c.paragraphs[0], val, bold=bold, size=9,
                 color="0F172A" if bold else "1E293B", align=align)

    # ── PROJECTED INCOME ──────────────────────────────────────
    _section_header(doc, "PROJECTED INCOME  (at market rents)")
    inc_rows = []
    for i, s in enumerate(comp_summary):
        n       = unit_cnts[i] if i < len(unit_cnts) else 0
        monthly = n * (s.get("avg_rent") or 0)
        inc_rows.append(
            (f"  {s['beds']}BD/{s['baths']}BA  ×  {n} units",
             f"${monthly:,.0f}/mo", False)
        )
    inc_rows.append(("Est. Gross Monthly Revenue", f"${gross_monthly:,.0f}",           True))
    inc_rows.append(("Est. Gross Annual Revenue",  f"${gross_annual:,.0f}",            True))
    inc_rows.append(("Unlevered CoC (Yr 1)",
                     f"{_unlev_coc:.1%}" if _all_in  else "—",               True))
    inc_rows.append(("Levered CoC (Yr 1)",
                     f"{_lev_coc:.1%}"   if _equity  else "—",               True))

    itbl = doc.add_table(rows=len(inc_rows), cols=2)
    itbl.width = CW
    _light_borders(itbl)

    divider = len(comp_summary)
    for i, (lbl, val, is_total) in enumerate(inc_rows):
        bg = BLUE_BG if is_total else (SLATE_BG if i % 2 == 0 else "FFFFFF")
        lc, vc = itbl.rows[i].cells[0], itbl.rows[i].cells[1]
        lc.width = Inches(4.5);  vc.width = Inches(2.5)
        _shading(lc, bg);        _shading(vc, bg)
        _cell_margins(lc);       _cell_margins(vc)
        col = "0F172A" if is_total else "475569"
        _run(lc.paragraphs[0], lbl, bold=is_total, size=9, color=col)
        _run(vc.paragraphs[0], val, bold=is_total, size=9, color=col,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── FOOTER NOTE ───────────────────────────────────────────
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(10)

    fn_p = doc.add_paragraph()
    fn_p.paragraph_format.space_before = Pt(0)
    fn_p.paragraph_format.space_after  = Pt(4)
    _para_border_bottom(fn_p, color="CBD5E1", sz="4")
    fn_p.paragraph_format.space_before = Pt(8)
    _run(fn_p,
         "The attached Excel model includes the full 10-year pro forma, debt schedule, "
         "returns analysis, and pricing scenarios — all populated with live market comp data. "
         "Blue cells are editable inputs.",
         size=8, color="64748B", italic=True)

    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.paragraph_format.space_before = Pt(6)
    _run(pf, f"Ping Underwriting Engine  ·  {date_str}", size=7.5, color="94A3B8")

    # ── PAGE 2: Market Analysis  /  PAGE 3+: Full comp listings ──
    if all_comp_rows:
        import statistics as _stats
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE

        # ── PAGE BREAK → Market Analysis ──────────────────────
        pb = doc.add_paragraph()
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after  = Pt(0)
        pb.add_run().add_break(__import__("docx").enum.text.WD_BREAK.PAGE)

        # ── MARKET ANALYSIS HEADER ─────────────────────────────
        ma_hdr = doc.add_table(rows=1, cols=2)
        ma_hdr.width = CW
        _tbl_borders(ma_hdr)
        mal = ma_hdr.rows[0].cells[0]
        mar = ma_hdr.rows[0].cells[1]
        mal.width = Inches(5.0)
        mar.width = Inches(2.0)
        for _c in (mal, mar):
            _shading(_c, NAVY)
            _cell_margins(_c, top=80, bottom=80, left=160, right=160)
        _run(mal.paragraphs[0], "MARKET ANALYSIS",
             bold=True, size=10, color="FFFFFF")
        _run(mar.paragraphs[0],
             f"{len(all_comp_rows)} comps  ·  {search_meta.get('radius')} mi radius",
             size=8, color="94A3B8", align=WD_ALIGN_PARAGRAPH.RIGHT)

        # ── Pre-compute stats ──────────────────────────────────
        from collections import OrderedDict as _OD

        def _percentile(sorted_vals, pct):
            if not sorted_vals:
                return 0
            idx = (len(sorted_vals) - 1) * pct / 100
            lo  = int(idx)
            hi  = min(lo + 1, len(sorted_vals) - 1)
            return sorted_vals[lo] + (sorted_vals[hi] - sorted_vals[lo]) * (idx - lo)

        unit_groups = _OD()
        for _row in all_comp_rows:
            _key = (_row["filter_beds"], _row["filter_baths"])
            unit_groups.setdefault(_key, []).append(_row)

        total_comp_n = len(all_comp_rows)

        # ── SECTION: Rent Distribution ─────────────────────────
        _section_header(doc, "RENT DISTRIBUTION BY UNIT TYPE")

        rd_hdrs   = ["Unit Type", "n", "Min", "25th %ile",
                     "Median", "Avg", "75th %ile", "Max", "Std Dev"]
        rd_widths = [Inches(1.05), Inches(0.40), Inches(0.75), Inches(0.80),
                     Inches(0.80), Inches(0.80), Inches(0.80), Inches(0.75),
                     Inches(0.85)]

        rdtbl = doc.add_table(rows=1 + len(unit_groups), cols=len(rd_hdrs))
        rdtbl.width = CW
        _light_borders(rdtbl)

        for _ci, (_h, _w) in enumerate(zip(rd_hdrs, rd_widths)):
            _c = rdtbl.rows[0].cells[_ci]
            _c.width = _w
            _shading(_c, LIGHT_NAVY)
            _cell_margins(_c, top=40, bottom=40, left=60, right=60)
            _p = _c.paragraphs[0]
            _p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if _ci == 0
                            else WD_ALIGN_PARAGRAPH.CENTER)
            _run(_p, _h, bold=True, size=7.5, color="FFFFFF")

        for _ri, ((_beds, _baths), _grp) in enumerate(unit_groups.items()):
            _bg    = SLATE_BG if _ri % 2 == 0 else "FFFFFF"
            _rents = sorted([_r["price"] for _r in _grp if _r.get("price")])
            try:
                _bi  = int(float(_beds))
                _baf = float(_baths)
                _bas = str(int(_baf)) if _baf == int(_baf) else str(_baf)
                _ulabel = f"{_bi}BD/{_bas}BA"
            except Exception:
                _ulabel = f"{_beds}BD/{_baths}BA"
            if _rents:
                _n  = len(_rents)
                _sd = _stats.stdev(_rents) if _n > 1 else 0
                _rd_vals = [
                    (_ulabel,                            WD_ALIGN_PARAGRAPH.LEFT,   False),
                    (str(_n),                            WD_ALIGN_PARAGRAPH.CENTER, False),
                    (f"${_rents[0]:,.0f}",               WD_ALIGN_PARAGRAPH.CENTER, False),
                    (f"${_percentile(_rents,25):,.0f}",  WD_ALIGN_PARAGRAPH.CENTER, False),
                    (f"${_percentile(_rents,50):,.0f}",  WD_ALIGN_PARAGRAPH.CENTER, True),
                    (f"${_stats.mean(_rents):,.0f}",     WD_ALIGN_PARAGRAPH.CENTER, True),
                    (f"${_percentile(_rents,75):,.0f}",  WD_ALIGN_PARAGRAPH.CENTER, False),
                    (f"${_rents[-1]:,.0f}",              WD_ALIGN_PARAGRAPH.CENTER, False),
                    (f"${_sd:,.0f}",                     WD_ALIGN_PARAGRAPH.CENTER, False),
                ]
            else:
                _rd_vals = [(_ulabel, WD_ALIGN_PARAGRAPH.LEFT, False)] + \
                           [("—", WD_ALIGN_PARAGRAPH.CENTER, False)] * 8
            _drow = rdtbl.rows[_ri + 1]
            for _ci, (_val, _aln, _bld) in enumerate(_rd_vals):
                _c = _drow.cells[_ci]
                _c.width = rd_widths[_ci]
                _shading(_c, _bg)
                _cell_margins(_c, top=40, bottom=40, left=60, right=60)
                _p = _c.paragraphs[0]
                _p.alignment = _aln
                _run(_p, _val, bold=_bld, size=8,
                     color="0F172A" if _bld else "374151")

        # ── SECTION: Property Type Mix ─────────────────────────
        _section_header(doc, "PROPERTY TYPE MIX")

        _type_counts = defaultdict(int)
        for _row in all_comp_rows:
            _pt = ((_row.get("propertyType") or "Unknown")
                   .replace("SingleFamily", "SFR"))
            _type_counts[_pt] += 1
        _pt_sorted = sorted(_type_counts.items(), key=lambda _x: -_x[1])

        pt_hdrs   = ["Property Type", "# Comps", "Share",
                     "Avg Rent/Mo", "Avg SF", "Avg Rent/SF"]
        pt_widths = [Inches(1.60), Inches(0.90), Inches(0.80),
                     Inches(1.20), Inches(1.20), Inches(1.30)]

        pttbl = doc.add_table(rows=1 + len(_pt_sorted), cols=len(pt_hdrs))
        pttbl.width = CW
        _light_borders(pttbl)

        for _ci, (_h, _w) in enumerate(zip(pt_hdrs, pt_widths)):
            _c = pttbl.rows[0].cells[_ci]
            _c.width = _w
            _shading(_c, LIGHT_NAVY)
            _cell_margins(_c, top=40, bottom=40, left=60, right=60)
            _p = _c.paragraphs[0]
            _p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if _ci == 0
                            else WD_ALIGN_PARAGRAPH.CENTER)
            _run(_p, _h, bold=True, size=7.5, color="FFFFFF")

        for _ri, (_ptype, _cnt) in enumerate(_pt_sorted):
            _bg    = SLATE_BG if _ri % 2 == 0 else "FFFFFF"
            _share = _cnt / total_comp_n * 100 if total_comp_n else 0
            _pt_rows  = [_r for _r in all_comp_rows
                         if ((_r.get("propertyType") or "")
                             .replace("SingleFamily", "SFR")) == _ptype]
            _pt_rents = [_r["price"]        for _r in _pt_rows if _r.get("price")]
            _pt_sqfts = [_r["squareFootage"] for _r in _pt_rows
                         if _r.get("squareFootage")]
            _pt_avg_r = _stats.mean(_pt_rents) if _pt_rents else 0
            _pt_avg_s = _stats.mean(_pt_sqfts) if _pt_sqfts else 0
            _pt_rsf   = _pt_avg_r / _pt_avg_s if _pt_avg_s else 0
            _pt_vals  = [
                (_ptype,                          WD_ALIGN_PARAGRAPH.LEFT,   False),
                (str(_cnt),                       WD_ALIGN_PARAGRAPH.CENTER, False),
                (f"{_share:.0f}%",                WD_ALIGN_PARAGRAPH.CENTER, False),
                (f"${_pt_avg_r:,.0f}" if _pt_avg_r else "—",
                                                  WD_ALIGN_PARAGRAPH.CENTER, True),
                (f"{_pt_avg_s:,.0f}" if _pt_avg_s else "—",
                                                  WD_ALIGN_PARAGRAPH.CENTER, False),
                (f"${_pt_rsf:.2f}" if _pt_rsf else "—",
                                                  WD_ALIGN_PARAGRAPH.CENTER, True),
            ]
            _drow = pttbl.rows[_ri + 1]
            for _ci, (_val, _aln, _bld) in enumerate(_pt_vals):
                _c = _drow.cells[_ci]
                _c.width = pt_widths[_ci]
                _shading(_c, _bg)
                _cell_margins(_c, top=40, bottom=40, left=60, right=60)
                _p = _c.paragraphs[0]
                _p.alignment = _aln
                _run(_p, _val, bold=_bld, size=8,
                     color="0F172A" if _bld else "374151")

        # ── SECTION: Days on Market ────────────────────────────
        _doms = []
        for _row in all_comp_rows:
            try:
                _d = _row.get("daysOnMarket")
                if _d not in (None, "", "—"):
                    _doms.append(int(_d))
            except (TypeError, ValueError):
                pass
        if _doms:
            _section_header(doc, "DAYS ON MARKET")
            dom_tbl = doc.add_table(rows=2, cols=4)
            dom_tbl.width = CW
            _light_borders(dom_tbl)
            _dom_stats = [
                ("Average",  f"{_stats.mean(_doms):.1f} days"),
                ("Median",   f"{_stats.median(_doms):.0f} days"),
                ("Fastest",  f"{min(_doms)} days"),
                ("Slowest",  f"{max(_doms)} days"),
            ]
            _dom_w = [Inches(1.75)] * 4
            for _ci, (_lbl, _val) in enumerate(_dom_stats):
                _hc = dom_tbl.rows[0].cells[_ci]
                _dc = dom_tbl.rows[1].cells[_ci]
                _hc.width = _dom_w[_ci]
                _dc.width = _dom_w[_ci]
                _shading(_hc, LIGHT_NAVY)
                _shading(_dc, BLUE_BG)
                _cell_margins(_hc, top=50, bottom=50, left=80, right=80)
                _cell_margins(_dc, top=60, bottom=60, left=80, right=80)
                _ph = _hc.paragraphs[0]
                _ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(_ph, _lbl, bold=True, size=8, color="FFFFFF")
                _pd = _dc.paragraphs[0]
                _pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(_pd, _val, bold=True, size=11, color="1E3A5F")

        # ── SECTION: Distance Distribution ─────────────────────
        _section_header(doc, "DISTANCE DISTRIBUTION")

        _dist_bins   = [(0, 0.5), (0.5, 1.0), (1.0, 1.5), (1.5, 99)]
        _dist_labels = ["< 0.5 km", "0.5 – 1.0 km", "1.0 – 1.5 km", "> 1.5 km"]
        _dist_counts = [
            sum(1 for _r in all_comp_rows
                if _lo <= _r.get("distance_km", 0) < _hi)
            for (_lo, _hi) in _dist_bins
        ]

        dd_hdrs   = ["Distance Band", "# Comps", "Share", "Avg Rent/Mo", "Avg SF"]
        dd_widths = [Inches(1.60), Inches(0.90), Inches(0.80),
                     Inches(1.90), Inches(1.80)]

        ddtbl = doc.add_table(rows=1 + len(_dist_bins), cols=len(dd_hdrs))
        ddtbl.width = CW
        _light_borders(ddtbl)

        for _ci, (_h, _w) in enumerate(zip(dd_hdrs, dd_widths)):
            _c = ddtbl.rows[0].cells[_ci]
            _c.width = _w
            _shading(_c, LIGHT_NAVY)
            _cell_margins(_c, top=40, bottom=40, left=60, right=60)
            _p = _c.paragraphs[0]
            _p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if _ci == 0
                            else WD_ALIGN_PARAGRAPH.CENTER)
            _run(_p, _h, bold=True, size=7.5, color="FFFFFF")

        for _ri, ((_lo, _hi), _lbl, _cnt) in enumerate(
                zip(_dist_bins, _dist_labels, _dist_counts)):
            _bg    = SLATE_BG if _ri % 2 == 0 else "FFFFFF"
            _share = _cnt / total_comp_n * 100 if total_comp_n else 0
            _band  = [_r for _r in all_comp_rows
                      if _lo <= _r.get("distance_km", 0) < _hi]
            _b_rents = [_r["price"]        for _r in _band if _r.get("price")]
            _b_sqfts = [_r["squareFootage"] for _r in _band
                        if _r.get("squareFootage")]
            _avg_r = _stats.mean(_b_rents) if _b_rents else 0
            _avg_s = _stats.mean(_b_sqfts) if _b_sqfts else 0
            _dd_vals = [
                (_lbl,                     WD_ALIGN_PARAGRAPH.LEFT,   False),
                (str(_cnt),                WD_ALIGN_PARAGRAPH.CENTER, False),
                (f"{_share:.0f}%",         WD_ALIGN_PARAGRAPH.CENTER, False),
                (f"${_avg_r:,.0f}" if _avg_r else "—",
                                           WD_ALIGN_PARAGRAPH.CENTER, True),
                (f"{_avg_s:,.0f} SF" if _avg_s else "—",
                                           WD_ALIGN_PARAGRAPH.CENTER, False),
            ]
            _drow = ddtbl.rows[_ri + 1]
            for _ci, (_val, _aln, _bld) in enumerate(_dd_vals):
                _c = _drow.cells[_ci]
                _c.width = dd_widths[_ci]
                _shading(_c, _bg)
                _cell_margins(_c, top=40, bottom=40, left=60, right=60)
                _p = _c.paragraphs[0]
                _p.alignment = _aln
                _run(_p, _val, bold=_bld, size=8,
                     color="0F172A" if _bld else "374151")

        # Market Analysis page footer
        ma_foot = doc.add_paragraph()
        ma_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ma_foot.paragraph_format.space_before = Pt(8)
        _run(ma_foot,
             f"Ping Underwriting Engine  ·  {search_meta.get('searchId')}  ·  {date_str}",
             size=7.5, color="94A3B8")

        # ── PAGE BREAK → Comp Listings ────────────────────────
        pb2 = doc.add_paragraph()
        pb2.paragraph_format.space_before = Pt(0)
        pb2.paragraph_format.space_after  = Pt(0)
        pb2.add_run().add_break(__import__("docx").enum.text.WD_BREAK.PAGE)

        # Comp listings header — slim navy bar
        p2hdr = doc.add_table(rows=1, cols=2)
        p2hdr.width = CW
        _tbl_borders(p2hdr)
        p2l = p2hdr.rows[0].cells[0]
        p2r = p2hdr.rows[0].cells[1]
        p2l.width = Inches(5.0)
        p2r.width = Inches(2.0)
        for c in (p2l, p2r):
            _shading(c, NAVY)
            _cell_margins(c, top=80, bottom=80, left=160, right=160)
        _run(p2l.paragraphs[0],
             "COMPARABLE RENTAL LISTINGS",
             bold=True, size=10, color="FFFFFF")
        _run(p2r.paragraphs[0],
             f"{len(all_comp_rows)} total comps  ·  {search_meta.get('radius')} mi radius",
             size=8, color="94A3B8",
             align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Group rows by (filter_beds, filter_baths) — preserve insertion order
        from collections import OrderedDict
        groups = OrderedDict()
        for row in all_comp_rows:
            key = (row["filter_beds"], row["filter_baths"])
            groups.setdefault(key, []).append(row)

        # Column spec: (header, width_inches, field_fn, align, bold)
        COMP_COLS = [
            ("#",          0.35, lambda r: r["rank"],                                   WD_ALIGN_PARAGRAPH.CENTER, False),
            ("Address",    2.65, lambda r: r["formattedAddress"],                        WD_ALIGN_PARAGRAPH.LEFT,   False),
            ("Rent/Mo",    0.80, lambda r: f"${r['price']:,.0f}" if r["price"] else "—", WD_ALIGN_PARAGRAPH.CENTER, True),
            ("Bd/Ba",      0.55, lambda r: f"{int(r['filter_beds'])}bd/{r['filter_baths']:.0f}ba" if r["filter_beds"] else "—",
                                                                                          WD_ALIGN_PARAGRAPH.CENTER, False),
            ("SF",         0.65, lambda r: f"{int(r['squareFootage']):,}" if r["squareFootage"] else "—",
                                                                                          WD_ALIGN_PARAGRAPH.CENTER, False),
            ("$/SF",       0.65, lambda r: f"${r['price']/r['squareFootage']:.2f}" if r["price"] and r["squareFootage"] else "—",
                                                                                          WD_ALIGN_PARAGRAPH.CENTER, True),
            ("Dist (km)",  0.60, lambda r: f"{r['distance_km']:.2f}",                    WD_ALIGN_PARAGRAPH.CENTER, False),
            ("Type",       0.75, lambda r: (r["propertyType"] or "").replace("SingleFamily", "SFR")[:14],
                                                                                          WD_ALIGN_PARAGRAPH.LEFT,   False),
        ]
        col_widths = [Inches(c[1]) for c in COMP_COLS]

        for (beds, baths), rows in groups.items():
            # Unit type sub-header
            try:
                b_int = int(float(beds))
                ba    = float(baths)
                ba_s  = str(int(ba)) if ba == int(ba) else str(ba)
                type_label = f"{b_int}BD / {ba_s}BA — {len(rows)} comps"
            except Exception:
                type_label = f"{beds}BD / {baths}BA — {len(rows)} comps"

            sh = doc.add_paragraph()
            sh.paragraph_format.space_before = Pt(10)
            sh.paragraph_format.space_after  = Pt(3)
            _run(sh, type_label, bold=True, size=9, color="1E3A5F")
            _para_border_bottom(sh, color="3B82F6", sz="4")

            # Table for this group
            tbl = doc.add_table(rows=1 + len(rows), cols=len(COMP_COLS))
            tbl.width = CW
            _light_borders(tbl)

            # Header row
            hrow = tbl.rows[0]
            for ci, (hdr_txt, _, _, align, _bold) in enumerate(COMP_COLS):
                c = hrow.cells[ci]
                c.width = col_widths[ci]
                _shading(c, NAVY)
                _cell_margins(c, top=40, bottom=40, left=80, right=80)
                p = c.paragraphs[0]
                p.alignment = align
                _run(p, hdr_txt, bold=True, size=8, color="FFFFFF")

            # Data rows
            for ri, comp in enumerate(rows):
                drow = tbl.rows[ri + 1]
                bg   = SLATE_BG if ri % 2 == 0 else "FFFFFF"
                for ci, (_, _, field_fn, align, bold) in enumerate(COMP_COLS):
                    c = drow.cells[ci]
                    c.width = col_widths[ci]
                    _shading(c, bg)
                    _cell_margins(c, top=36, bottom=36, left=80, right=80)
                    try:
                        val = str(field_fn(comp))
                    except Exception:
                        val = "—"
                    p = c.paragraphs[0]
                    p.alignment = align
                    _run(p, val, bold=bold, size=8,
                         color="0F172A" if bold else "374151")

        # Page 2 footer
        p2f = doc.add_paragraph()
        p2f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2f.paragraph_format.space_before = Pt(8)
        _run(p2f,
             f"Ping Underwriting Engine  ·  {search_meta.get('searchId')}  ·  {date_str}",
             size=7.5, color="94A3B8")

    # ── Save + patch settings.xml ─────────────────────────────
    import zipfile, re as _re
    out_path = job_dir / f"Ping_{search_meta['searchId']}_Summary.docx"
    doc.save(str(out_path))

    # python-docx omits the required w:percent on w:zoom — patch it.
    _bytes = out_path.read_bytes()
    with zipfile.ZipFile(out_path, "r") as z:
        names = z.namelist()
        if "word/settings.xml" in names:
            settings_xml = z.read("word/settings.xml").decode("utf-8")
            # Add percent="100" if the zoom element is missing it
            settings_xml = _re.sub(
                r'(<w:zoom\b(?![^>]*w:percent)[^>]*?)(\s*/>)',
                r'\1 w:percent="100"\2',
                settings_xml,
            )
            # Rewrite zip with patched settings.xml
            import io as _io
            buf = _io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
                with zipfile.ZipFile(out_path, "r") as zin:
                    for item in zin.infolist():
                        if item.filename == "word/settings.xml":
                            zout.writestr(item, settings_xml.encode("utf-8"))
                        else:
                            zout.writestr(item, zin.read(item.filename))
            out_path.write_bytes(buf.getvalue())

    print(f"         Summary docx: {out_path.name}")
    return out_path


# ── Email Body ────────────────────────────────────────────────────────────────